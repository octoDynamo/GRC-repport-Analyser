"""Document parser — supports PDF, DOCX, and XLSX extraction to text."""

from loguru import logger


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF using PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages = []
        for page_num, page in enumerate(doc, 1):
            text = page.get_text()
            if text.strip():
                pages.append(f"=== Page {page_num} ===\n{text}")
        doc.close()
        return "\n\n".join(pages)
    except Exception as exc:
        logger.error(f"PDF extraction failed: {exc}")
        raise


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file."""
    import io

    from docx import Document

    try:
        doc = Document(io.BytesIO(file_bytes))
        sections = []
        for para in doc.paragraphs:
            if para.text.strip():
                if para.style.name.startswith("Heading"):
                    sections.append(f"\n### {para.text.strip()}")
                else:
                    sections.append(para.text.strip())
        # Also extract table content
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    sections.append(row_text)
        return "\n".join(sections)
    except Exception as exc:
        logger.error(f"DOCX extraction failed: {exc}")
        raise


def extract_text_from_xlsx(file_bytes: bytes) -> str:
    """Extract text from an XLSX file (all sheets)."""
    import io

    import openpyxl

    try:
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
        sheets = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows_text = []
            for row in ws.iter_rows(values_only=True):
                row_values = [str(cell) for cell in row if cell is not None]
                if row_values:
                    rows_text.append(" | ".join(row_values))
            if rows_text:
                sheets.append(f"=== Sheet: {sheet_name} ===\n" + "\n".join(rows_text))
        return "\n\n".join(sheets)
    except Exception as exc:
        logger.error(f"XLSX extraction failed: {exc}")
        raise


def extract_text(file_bytes: bytes, file_format: str) -> str:
    """
    Main entry point for document text extraction.
    file_format should be: 'pdf', 'docx', or 'xlsx'
    """
    fmt = file_format.lower().strip(".")
    if fmt == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif fmt == "docx":
        return extract_text_from_docx(file_bytes)
    elif fmt == "xlsx":
        return extract_text_from_xlsx(file_bytes)
    else:
        raise ValueError(f"Unsupported file format: {file_format}")


def chunk_text(text: str, chunk_size: int = 1500, overlap: int = 200) -> list[str]:
    """
    Split text into overlapping chunks for RAG indexing.
    Returns a list of text chunks.
    """
    if len(text) <= chunk_size:
        return [text]

    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        chunks.append(chunk)
        if end >= len(text):
            break
        start = end - overlap
    return chunks
